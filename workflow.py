#TODO: 
# 1) Add a bot that connects to google api document generator and uses contact informatation
# 2) 

from openai import OpenAI
import os

import os
from textwrap import dedent
import docx
from docx import Document

from fuzzywuzzy import fuzz

def fuzzy_split(text, pattern, threshold=50):
    """Splits text based on fuzzy matching with a given pattern."""

    # Find the index with the highest similarity score above the threshold
    best_match_index = max(
        range(len(text)),
        key=lambda i: fuzz.partial_ratio(text[i : i + len(pattern)], pattern)
    )

    if fuzz.partial_ratio(text[best_match_index : best_match_index + len(pattern)], pattern) > threshold:
        return [
            text[:best_match_index],
            text[best_match_index + len(pattern) :],
        ]
    else:
        return [text]  # No match found above the threshold


client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
model_name = 'gpt-4o'

def load_document(file_path):
    if file_path.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format. Only TXT and DOCX are supported.")
    return text

job_listing = load_document("Materials/job_listing.txt")
resume = load_document("Materials/Robbins_resume.docx")

#####################################################################

# Your entire prompt for the agent
prompt = dedent(f'''

Populate a JSON file based on information from the provided job listing.

Job listing: {job_listing}

expected_output:
A JSON-like output with the following information:

Full role:
Company name:
The most important day-to-day challenge for this role:
Skill 1:
Skill 2:
Skill 3:
Skill 4:
Skill 5:
Skill 6:

Do not include explanations, reasoning, or additional commentary.
''')

# API call (using the Chat Completions API)
response = client.chat.completions.create(model=model_name,  # Or another suitable model
messages=[
    {"role": "system", "content": "You are an expert in talent acquisition and workforce optimization with 20 years of experience summarizing job descriptions"},
    {"role": "user", "content": prompt}
])
# Get the assistant's response and store it in a variable
job_summary = response.choices[0].message.content
print(job_summary)

###########################################################################

# Your entire prompt for the agent
prompt = dedent(f'''
Your task is to extract information from the resume and output in JSON format.

Extract the name and overall title of the applicant.
Also, extract ALL previous work experiences, including company name, job title,
and the word for word description of job duties found in the resume.

Resume: {resume}

expected_output:
A JSON file with the following format:

Applicant name:
Applicant title:
Applicant email:
Applicant residence:
Applicant portfolio website:
Applicant github:

Previous work experience:
[
{{
    "company": "",
    "title": "",
    "responsibilities": ""
}},
# ... more work experiences as needed
]

Do not include explanations, reasoning, or additional commentary.
''')

# API call (using the Chat Completions API)
response = client.chat.completions.create(model=model_name,  # Or another suitable model
messages=[
    {"role": "system", "content": "You are an expert in talent acquisition and workforce optimization with 20 years of experience summarizing job descriptions."},
    {"role": "user", "content": prompt}
])
# Get the assistant's response and store it in a variable
resume_summary = response.choices[0].message.content
print(resume_summary)

pattern = "Previous work experience"
split_parts = fuzzy_split(resume_summary, pattern, threshold=50)

if len(split_parts) == 2:  
    contact_info = split_parts[0].strip()
    resume_summary = split_parts[1].strip()
    print('########################### Contact info split successful#########################\n\n')
else:
    # Handle the case where the pattern isn't found or doesn't meet the threshold
    print("Pattern not found or similarity too low.")

###########################################################################

# Your entire prompt for the agent
prompt = dedent(f'''
Using the provided job summary and resume details, write a compelling opening paragraph (hook) for the cover letter. The hook should:
- Be less than 100 words.
- Explain why the applicant can succesfully take on the challenges of the job outlined in the job summary.
- Use keywords that will resonate with ATS scans.
- Incorporate total years experience from the applicant if it appears in the resume.

Resume details: {resume_summary}
Job summary: {job_summary}

expected_output: >
An attention-grabbing cover letter hook (less than 100 words).
The hook should start with:

I was thrilled to see your listing for [role mentioned above], because it is exactly the job I've been looking for."

Do not include explanations, reasoning, or additional commentary.
''')

# API call (using the Chat Completions API)
response = client.chat.completions.create(model=model_name,  # Or another suitable model
messages=[
    {"role": "system", "content": "You are an expert cover letter writer with a comprehensive understanding of Applicant Tracking Systems (ATS) and keyword optimization."},
    {"role": "user", "content": prompt}
])
# Get the assistant's response and store it in a variable
hook = response.choices[0].message.content
print(hook + '\n') 

###########################################################################

# Your entire prompt for the agent
prompt = dedent(f'''
Based on the provided hook, resume details, and job summary write a cover letter body with a strong focus on the company's future needs and how the applicant can fulfill those needs.

Take a deep breath and think through the writing process step by step.

- Output the body as TWO paragraphs.
- For each paragraph, pick ONE skill from the job summary and do the following:
1) Choose 2 relevant quantifieable metrics that directly address the ONE skill from the job summary.
2) Explicitely write why each quantifiable metric aligns with the chosen job summary skill.

- Ensure the entire length is around 200 words.
- Heavily prioritize unique descriptors and unconventional writing style.

Resume details: {resume_summary}
Job summary: {job_summary}
Hook: {hook}

expected_output:
2 paragraph cover letter body with a strong focus on the company's future needs and how the applicant can fulfill those needs.
DO NOT output the provided hook.
Do not include explanations, reasoning, or additional commentary.
''')

# API call (using the Chat Completions API)
response = client.chat.completions.create(model=model_name,  # Or another suitable model
messages=[
    {"role": "system", "content": "You are an expert cover letter writer with a comprehensive understanding of Applicant Tracking Systems (ATS) and keyword optimization."},
    {"role": "user", "content": prompt}
])
# Get the assistant's response and store it in a variable
body = response.choices[0].message.content
print(body + '\n\n')

###########################################################################

# Your entire prompt for the agent
prompt = dedent(f'''
Based on the provided cover letter hook and body, write a conclusion paragraph that summarizes why the applicant is a great fit for the job.

- Output the conclusion as ONE paragraph.
- Summarize why the applicant's strengths align with the job requirements.
- Include a call to action at the end of the paragraph.
- Ensure the entire length is under 75 words.
- Heavily prioritize unique descriptors and unconventional writing style.

Hook: {hook}
Body: {body}

expected_output:
1 paragraph cover letter conclusion.
DO NOT output the provided hook or body.
Do not include explanations, reasoning, or additional commentary.
''')

# API call (using the Chat Completions API)
response = client.chat.completions.create(model=model_name,  # Or another suitable model
messages=[
    {"role": "system", "content": "You are an expert cover letter writer with a comprehensive understanding of Applicant Tracking Systems (ATS) and keyword optimization."},
    {"role": "user", "content": prompt}
])
# Get the assistant's response and store it in a variable
conclusion = response.choices[0].message.content
print(conclusion + '\n\n')

###########################################################################

# Your entire prompt for the agent
prompt = dedent(f'''
Analyze the provided cover letter hook and body and compare them against the provided resume. 
Identify any statements in the cover letter that are inconsistent with, or not supported by, the details in the resume. 
Report any discrepancies found.

Hook: {hook}
Body: {body}
Conclusion : {conclusion}
Resume: {resume}

expected_output:
A detailed analysis on details within the cover letter that are not supported by the information in the resume.
''')

# API call (using the Chat Completions API)
response = client.chat.completions.create(model=model_name,  # Or another suitable model
messages=[
    {"role": "system", "content": "You are a meticulous, detail-oriented, and impartial editor that is not swayed by persuasive language but focuses solely on factual alignment."},
    {"role": "user", "content": prompt}
])
# Get the assistant's response and store it in a variable
truth_analysis = response.choices[0].message.content
print(truth_analysis + '\n\n')

###########################################################################

# Your entire prompt for the agent
prompt = dedent(f'''
Rewrite the provided cover letter using the comments from the fact analysis.
ONLY make changes to elements that are factually incorrect while preserving the original content as much as possible.

Resume: 
{hook}
{body}
{conclusion}
Fact analysis: {truth_analysis}

expected_output:
A full revised version of the cover letter with only factually incorrect elements corrected.
Do not include explanations, reasoning, or additional commentary.
''')

# API call (using the Chat Completions API)
response = client.chat.completions.create(model=model_name,  # Or another suitable model
messages=[
    {"role": "system", "content": "You are an expert cover letter writer with a comprehensive understanding of Applicant Tracking Systems (ATS) and keyword optimization."},
    {"role": "user", "content": prompt}
])
# Get the assistant's response and store it in a variable
full_cover_letter = response.choices[0].message.content
print(full_cover_letter)

def export_to_docx(content, filename="cover_letter.docx"):
    """Exports text content to a Word document (.docx)."""

    document = Document()
    document.add_paragraph(content)
    document.save(filename)

output_file_path = "Finished_cover_letters/cover_letter.docx"
content = f"Dear Hiring Manager,\n\n{full_cover_letter}\n\nSincerely\nColton Robbins"
export_to_docx(content, output_file_path)
